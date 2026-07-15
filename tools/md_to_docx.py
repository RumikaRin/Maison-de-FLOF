import os
import re
from pathlib import Path
from PIL import Image

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Root path of the workspace
ROOT = Path(__file__).resolve().parents[1]

# XML Helper Functions for Premium Styling
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))

def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

def clean_xml_string(s):
    if not s:
        return ""
    # Filter out characters not allowed in XML:
    # Allowed: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    return "".join(
        ch for ch in s
        if (ord(ch) in (0x9, 0xA, 0xD) or
            (0x20 <= ord(ch) <= 0xD7FF) or
            (0xE000 <= ord(ch) <= 0xFFFD) or
            (0x10000 <= ord(ch) <= 0x10FFFF))
    )

# Helper function to add formatting runs (bold, italic, code) to a paragraph
def add_runs_to_paragraph(p, text, font_name="Times New Roman", font_size=Pt(13), is_h1=False, is_h3=False, is_h4=False):
    # Regex to split on bold (**text**), italic (*text* or _text_), and inline code (`code`)
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", re.DOTALL)
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Word compatibility properties for fonts
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        
        if part.startswith("**") and part.endswith("**"):
            run.text = clean_xml_string(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = clean_xml_string(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = clean_xml_string(part[1:-1])
            run.font.name = "Consolas"
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), "Consolas")
            rFonts.set(qn('w:hAnsi'), "Consolas")
            rPr.append(rFonts)
        else:
            run.text = clean_xml_string(part)
            
        if is_h1:
            run.bold = True
        if is_h3:
            run.bold = True
            run.italic = True
        if is_h4:
            run.italic = True
            run.bold = False

def main():
    md_path = r"c:\Users\sansm\Downloads\BaoCao_FLOF_Full.md"
    docx_path = r"c:\Users\sansm\Downloads\BaoCao_FLOF_Full_with_Placeholders.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        return
        
    print("Reading Markdown report...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Configure page geometry (A4, Margins: Top 2cm, Bottom 3cm, Left 3cm, Right 2cm)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.5)
    
    # Configure document header/footer
    sec.header.paragraphs[0].text = "ĐỒ ÁN TỐT NGHIỆP - WEBSITE MAISON DE FLOF"
    sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in sec.header.paragraphs[0].runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(128, 128, 128)
        
    add_page_number(sec.footer.paragraphs[0])
    
    # Setup styles
    styles = doc.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Predefined Heading Styles customization
    for style_name, size, bold, italic, alignment in [
        ("Heading 1", 14, True, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, True, True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 4", 14, False, True, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.alignment = alignment
        style.paragraph_format.keep_with_next = True
        
    # Variables for state machine parsing
    in_cover = True
    cover_lines = []
    
    in_code_block = False
    code_lines = []
    
    in_table = False
    table_lines = []
    
    img_no = 0
    
    idx = 0
    total_lines = len(lines)
    
    while idx < total_lines:
        line = lines[idx]
        stripped = line.strip()
        
        # 1. Handle Cover Page
        if in_cover:
            if stripped == "---":
                in_cover = False
                # Format and output cover page
                print("Rendering cover page...")
                for cl in cover_lines:
                    cl_stripped = cl.strip()
                    if not cl_stripped:
                        continue
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Style based on content type
                    if "BỘ GIÁO DỤC" in cl_stripped or "TRƯỜNG ĐẠI HỌC" in cl_stripped:
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(10)
                        add_runs_to_paragraph(p, cl_stripped, font_size=Pt(14))
                        p.runs[0].bold = True
                    elif "Đồ án:" in cl_stripped or "Thiết kế và xây dựng" in cl_stripped or "Đề Tài:" in cl_stripped:
                        p.paragraph_format.space_before = Pt(24)
                        p.paragraph_format.space_after = Pt(24)
                        # Extract title text
                        title_text = cl_stripped.replace("<", "").replace(">", "").strip()
                        add_runs_to_paragraph(p, title_text, font_size=Pt(18))
                        p.runs[0].bold = True
                    elif "Tên:" in cl_stripped or "Lớp:" in cl_stripped or "Nhóm" in cl_stripped:
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(12)
                        add_runs_to_paragraph(p, cl_stripped, font_size=Pt(13))
                        p.runs[0].bold = True
                    elif "Hà Nội," in cl_stripped:
                        p.paragraph_format.space_before = Pt(60)
                        add_runs_to_paragraph(p, cl_stripped, font_size=Pt(13))
                        p.runs[0].italic = True
                    else:
                        add_runs_to_paragraph(p, cl_stripped, font_size=Pt(13))
                        
                doc.add_page_break()
                idx += 1
                continue
            else:
                cover_lines.append(line)
                idx += 1
                continue
                
        # 2. Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                in_code_block = False
                code_text = "".join(code_lines)
                
                # Check if this is a Mermaid diagram
                # We skip compiling raw mermaid syntax into text blocks, instead we let the images show it
                # but if user wants it, we can output it. Let's output it as a code block.
                # Actually, all diagrams already have image counterparts in the report.
                # To look professional, let's compile code text into a premium single-cell table.
                
                # Setup table for code block
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                # Width 6.2 inches is the printable width
                cell = table.cell(0, 0)
                cell.width = Inches(6.2)
                set_cell_shading(cell, "F8F9FA")
                set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
                
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                
                run = p.add_run(clean_xml_string(code_text.strip()))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(40, 44, 52)
                
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), "Consolas")
                rFonts.set(qn('w:hAnsi'), "Consolas")
                rPr.append(rFonts)
                
                # Space after the code table
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(6)
                
                code_lines = []
            else:
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue
            
        # 3. Handle Tables
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            idx += 1
            continue
        elif in_table:
            # We reached the end of the table
            in_table = False
            
            # Parse table lines
            # First line: headers
            # Second line: separator (e.g., |---|---|)
            # Remaining lines: data rows
            if len(table_lines) >= 2:
                # Extract headers
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                num_cols = len(headers)
                
                # Setup Word Table
                table = doc.add_table(rows=1, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                hdr_cells = table.rows[0].cells
                for i, h_text in enumerate(headers):
                    hdr_cells[i].text = clean_xml_string(h_text)
                    
                # Add data rows
                for r_line in table_lines[2:]:
                    cells_data = [c.strip() for c in r_line.split("|")[1:-1]]
                    # Pad cells if necessary
                    while len(cells_data) < num_cols:
                        cells_data.append("")
                    cells_data = cells_data[:num_cols]
                    
                    row_cells = table.add_row().cells
                    for i, val in enumerate(cells_data):
                        row_cells[i].text = clean_xml_string(val)
                
                # Style the table
                table.style = 'Table Grid'
                
                # Distribute widths evenly based on column count
                total_width_in = 6.2
                col_width = total_width_in / num_cols
                
                for row_idx, row in enumerate(table.rows):
                    prevent_row_split(row)
                    if row_idx == 0:
                        set_repeat_table_header(row)
                    for col_idx, cell in enumerate(row.cells):
                        cell.width = Inches(col_width)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        text = p.text
                        p.text = ""
                        add_runs_to_paragraph(p, text, font_size=Pt(12))
                        
                        if row_idx == 0:
                            set_cell_shading(cell, "F2F4F7")
                            for run in p.runs:
                                run.bold = True
                                
                # Space after table
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(6)
                
            table_lines = []
            # Do not increment idx here, as the current line needs to be parsed in the next iteration
            continue
            
        # 4. Handle Empty lines
        if not stripped:
            idx += 1
            continue
            
        # 5. Handle Headings
        # Heading 1: starts with "# "
        # We also check if a Heading 2 starts with "## " but is an academic section (TOC, Preamble, Intro)
        # e.g., "## LỜI CẢM ƠN / LỜI NÓI ĐẦU", "## DANH MỤC TỪ VIẾT TẮT", "## MỞ ĐẦU BÁO CÁO", "## MỤC LỤC"
        # We should treat them as Heading 1 (centered, bold, 14pt, uppercase).
        is_h1_section = (
            stripped.startswith("## LỜI CẢM ƠN") or
            stripped.startswith("## MỤC LỤC") or
            stripped.startswith("## DANH MỤC TỪ VIẾT TẮT") or
            stripped.startswith("## MỞ ĐẦU BÁO CÁO")
        )
        
        if stripped.startswith("# ") or is_h1_section:
            # Heading 1
            title_text = stripped.lstrip("#").strip()
            # Convert to UPPERCASE
            title_text = title_text.upper()
            
            p = doc.add_paragraph(style="Heading 1")
            add_runs_to_paragraph(p, title_text, font_size=Pt(14), is_h1=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            
            idx += 1
            continue
            
        elif stripped.startswith("## "):
            # Heading 2
            title_text = stripped.lstrip("#").strip()
            
            p = doc.add_paragraph(style="Heading 2")
            add_runs_to_paragraph(p, title_text, font_size=Pt(14))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            
            idx += 1
            continue
            
        elif stripped.startswith("### "):
            # Heading 3
            title_text = stripped.lstrip("#").strip()
            
            p = doc.add_paragraph(style="Heading 3")
            add_runs_to_paragraph(p, title_text, font_size=Pt(14), is_h3=True)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            
            idx += 1
            continue
            
        elif stripped.startswith("#### "):
            # Heading 4
            title_text = stripped.lstrip("#").strip()
            
            p = doc.add_paragraph(style="Heading 4")
            add_runs_to_paragraph(p, title_text, font_size=Pt(14), is_h4=True)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            
            idx += 1
            continue
            
        # 6. Handle Images
        # e.g., ![Biểu đồ Use Case tổng quát](file:///d:/ProjectZ/FLOF/public/use_case_diagram.png)
        img_match = re.search(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img_match:
            caption = img_match.group(1).strip()
            img_url = img_match.group(2).strip()
            
            # Convert URL to local path
            clean_path = img_url.replace("file:///", "").replace("/", "\\")
            if clean_path.startswith("public") or clean_path.startswith("\\public"):
                clean_path = os.path.join(str(ROOT), clean_path.lstrip("\\/"))
            elif clean_path.startswith("d:\\ProjectZ\\FLOF\\public\\") or clean_path.startswith("d:/ProjectZ/FLOF/public/"):
                # Path is already absolute local
                pass
                
            if os.path.exists(clean_path):
                img_no += 1
                print(f"Adding image: {clean_path}")
                
                # Add picture paragraph and center it
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(6)
                
                # Check original image size using PIL and scale it if it exceeds printable width (6.2 inches)
                try:
                    with Image.open(clean_path) as img:
                        w_px, h_px = img.size
                        # Convert pixels to inches (assuming 96 dpi default if no info)
                        dpi = img.info.get("dpi", (96, 96))[0]
                        w_in = w_px / dpi
                        
                        if w_in > 6.2:
                            p_img.add_run().add_picture(clean_path, width=Inches(6.2))
                        else:
                            p_img.add_run().add_picture(clean_path, width=Inches(w_in))
                except Exception as e:
                    # Fallback to standard 6.0 width
                    print(f"  Warning scaling image: {e}")
                    p_img.add_run().add_picture(clean_path, width=Inches(6.0))
                    
                # Add centered caption below
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(4)
                p_cap.paragraph_format.space_after = Pt(12)
                p_cap.paragraph_format.keep_with_next = False
                
                cap_run = p_cap.add_run(f"Hình {img_no}. {caption}")
                cap_run.font.name = "Times New Roman"
                cap_run.font.size = Pt(12)
                cap_run.italic = True
            else:
                print(f"Warning: Image file not found at {clean_path}")
                
            idx += 1
            continue
            
        # 6b. Handle Screenshot Placeholders / Comments
        if "[CHÚ THÍCH THÊM ẢNH:" in stripped:
            # Extract the placeholder text
            placeholder_text = stripped.replace(">", "").replace("**", "").replace("[", "").replace("]", "").strip()
            
            # Add a premium callout box for placeholder
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            
            cell = table.cell(0, 0)
            cell.width = Inches(6.2)
            set_cell_shading(cell, "FFF3CD") # Light orange/yellow warning background
            set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f"【 {placeholder_text} 】")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(133, 100, 4) # Dark brown/gold warning color
            
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), "Times New Roman")
            rFonts.set(qn('w:hAnsi'), "Times New Roman")
            rPr.append(rFonts)
            
            # Space after the callout table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
            
            idx += 1
            continue
            
        # 7. Handle List Items
        # Bullet list item: starts with "* " or "- "
        is_bullet = stripped.startswith("* ") or stripped.startswith("- ")
        is_numbered = re.match(r"^\d+\.\s+", stripped)
        
        if is_bullet or is_numbered:
            if is_bullet:
                list_style = "List Bullet"
                list_text = stripped[2:].strip()
            else:
                list_style = "List Number"
                # Strip digit prefix and dot
                match_len = len(is_numbered.group(0))
                list_text = stripped[match_len:].strip()
                
            p = doc.add_paragraph(style=list_style)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            add_runs_to_paragraph(p, list_text, font_size=Pt(13))
            idx += 1
            continue
            
        # 8. Regular Body Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        add_runs_to_paragraph(p, stripped, font_size=Pt(13))
        idx += 1
        
    print(f"Saving compiled Word document to {docx_path}...")
    doc.save(docx_path)
    print("Document compiled successfully!")

if __name__ == "__main__":
    main()
