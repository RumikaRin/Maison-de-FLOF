from pathlib import Path
from datetime import date
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"
ASSETS = OUT / "assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

BLUE = "#2E74B5"
DARK = "#17365D"
INK = "#1F2937"
MUTED = "#64748B"
LIGHT = "#F2F4F7"
PALE = "#E8EEF5"
GREEN = "#0F766E"
GOLD = "#A16207"
RED = "#9B1C1C"
WHITE = "#FFFFFF"


def rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def arrow(draw, start, end, color=DARK, width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 15
    points = [
        (x2, y2),
        (x2 - length * math.cos(angle - 0.5), y2 - length * math.sin(angle - 0.5)),
        (x2 - length * math.cos(angle + 0.5), y2 - length * math.sin(angle + 0.5)),
    ]
    draw.polygon(points, fill=color)


def box(draw, xy, title, lines=(), fill=WHITE, outline=BLUE, title_fill=None, radius=16):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    if title_fill:
        draw.rounded_rectangle((x1, y1, x2, y1 + 48), radius=radius, fill=title_fill, outline=title_fill)
        draw.rectangle((x1, y1 + 25, x2, y1 + 48), fill=title_fill)
        title_color = WHITE
    else:
        title_color = DARK
    draw.text((x1 + 16, y1 + 12), title, font=font(24, True), fill=title_color)
    y = y1 + 60
    for line in lines:
        for piece in wrap(line, max(18, int((x2 - x1) / 13))):
            draw.text((x1 + 16, y), "• " + piece, font=font(18), fill=INK)
            y += 25


def save_canvas(name, size=(1500, 900)):
    image = Image.new("RGB", size, WHITE)
    return image, ImageDraw.Draw(image), ASSETS / name


def diagram_architecture():
    image, d, path = save_canvas("academic_architecture.png")
    d.text((55, 30), "Kiến trúc logic Maison de FLOF", font=font(38, True), fill=DARK)
    box(d, (70, 130, 370, 330), "Kênh người dùng", ["Storefront", "Tài khoản khách hàng", "Cổng quản trị"], PALE, BLUE, BLUE)
    box(d, (470, 105, 1020, 365), "Next.js Application", ["App Router + React 19", "Server/API routes", "Auth.js, Zod, RBAC", "Commerce services, audit"], LIGHT, BLUE, BLUE)
    box(d, (1120, 130, 1430, 330), "Dịch vụ ngoài", ["Cloudinary", "Resend", "MapLibre"], PALE, BLUE, BLUE)
    box(d, (310, 520, 750, 770), "Dữ liệu giao dịch", ["PostgreSQL / Neon", "Prisma ORM + migrations", "Order, Payment, Inventory", "AuditLog, Idempotency"], LIGHT, GREEN, GREEN)
    box(d, (850, 520, 1290, 770), "Vận hành & chất lượng", ["GitHub Actions CI", "Lint, typecheck, test, build", "Dependency audit mức high", "Cần bổ sung staging/monitoring"], LIGHT, GOLD, GOLD)
    arrow(d, (370, 230), (470, 230))
    arrow(d, (1020, 230), (1120, 230))
    arrow(d, (700, 365), (570, 520))
    arrow(d, (800, 365), (1030, 520))
    image.save(path)
    return path


def diagram_usecase():
    image, d, path = save_canvas("academic_usecase.png")
    d.text((55, 25), "Use Case tổng quan", font=font(38, True), fill=DARK)
    d.ellipse((45, 210, 145, 310), fill=PALE, outline=BLUE, width=3)
    d.text((44, 325), "Khách hàng", font=font(24, True), fill=DARK)
    d.ellipse((1340, 210, 1440, 310), fill=PALE, outline=BLUE, width=3)
    d.text((1325, 325), "Staff/Admin", font=font(24, True), fill=DARK)
    cases_left = ["Tìm kiếm & xem sơn/màu", "Tính lượng sơn", "Quản lý giỏ hàng", "Checkout COD/chuyển khoản", "Theo dõi đơn hàng", "Yêu cầu tư vấn/chat"]
    cases_right = ["Quản lý đơn & payment", "Đối soát/hoàn tiền", "Quản lý catalog/kho", "Quản lý nội dung", "Hỗ trợ khách hàng", "Xem audit log"]
    for i, text in enumerate(cases_left):
        y = 95 + i * 115
        d.ellipse((260, y, 720, y + 72), fill=LIGHT, outline=BLUE, width=3)
        d.text((290, y + 20), text, font=font(20, True), fill=INK)
        d.line((145, 260, 260, y + 36), fill=MUTED, width=2)
    for i, text in enumerate(cases_right):
        y = 95 + i * 115
        d.ellipse((780, y, 1240, y + 72), fill=LIGHT, outline=GREEN, width=3)
        d.text((810, y + 20), text, font=font(20, True), fill=INK)
        d.line((1240, y + 36, 1340, 260), fill=MUTED, width=2)
    image.save(path)
    return path


def diagram_erd():
    image, d, path = save_canvas("academic_erd.png", (1600, 1050))
    d.text((55, 25), "ERD rút gọn theo miền nghiệp vụ", font=font(38, True), fill=DARK)
    boxes = {
        "User / Customer": (70, 130, 390, 330, ["Role", "Address", "Wishlist", "Review"]),
        "Paint / Catalog": (640, 100, 960, 340, ["Category", "Supplier", "PaintColorLink", "Promotion"]),
        "Order": (1210, 130, 1530, 370, ["OrderItem", "OrderStatusHistory", "Coupon", "Address snapshot"]),
        "Payment": (1210, 520, 1530, 730, ["PaymentStatus", "Reference", "Paid/Refund time"]),
        "Inventory": (640, 680, 960, 900, ["InventoryTransaction", "Type", "Quantity", "Reference"]),
        "Governance": (70, 650, 390, 900, ["CheckoutIdempotency", "AuditLog", "Actor", "Metadata"]),
    }
    for title, (x1, y1, x2, y2, lines) in boxes.items():
        color = GREEN if title in ("Payment", "Inventory") else BLUE
        box(d, (x1, y1, x2, y2), title, lines, LIGHT, color, color)
    arrow(d, (390, 230), (640, 220))
    arrow(d, (960, 230), (1210, 245))
    arrow(d, (1370, 370), (1370, 520))
    arrow(d, (1210, 620), (960, 790))
    arrow(d, (640, 790), (390, 770))
    arrow(d, (250, 330), (250, 650))
    d.text((70, 970), "Ghi chú: sơ đồ nhóm 29 model thành các bounded context để dễ đọc.", font=font(20), fill=MUTED)
    image.save(path)
    return path


def diagram_activity():
    image, d, path = save_canvas("academic_activity_checkout.png", (1500, 1100))
    d.text((55, 25), "Activity Diagram: checkout có idempotency", font=font(38, True), fill=DARK)
    steps = [
        ("Khách gửi giỏ hàng + Idempotency-Key", BLUE),
        ("Xác thực người dùng và dữ liệu đầu vào", BLUE),
        ("Kiểm tra khóa đã xử lý?", GOLD),
        ("Tính giá phía server, kiểm tra coupon & tồn kho", BLUE),
        ("Transaction: tạo Order, Items, Payment, ledger", GREEN),
        ("Lưu kết quả idempotency", GREEN),
        ("Trả orderNumber và hướng dẫn thanh toán", BLUE),
    ]
    y = 105
    centers = []
    for text, color in steps:
        d.rounded_rectangle((350, y, 1150, y + 90), radius=25, fill=LIGHT, outline=color, width=4)
        d.text((390, y + 28), text, font=font(23, True), fill=INK)
        centers.append((750, y + 90))
        y += 135
    for i in range(len(centers) - 1):
        arrow(d, centers[i], (750, centers[i + 1][1] - 90))
    d.rounded_rectangle((40, 385, 285, 500), radius=20, fill=PALE, outline=GOLD, width=3)
    d.text((65, 412), "Có: trả lại kết quả\ncũ, không tạo đơn trùng", font=font(20, True), fill=INK)
    arrow(d, (350, 420), (285, 440), GOLD)
    image.save(path)
    return path


def diagram_sequence():
    image, d, path = save_canvas("academic_sequence_payment.png", (1700, 1000))
    d.text((55, 25), "Sequence Diagram: đối soát chuyển khoản thủ công", font=font(38, True), fill=DARK)
    actors = [("Khách hàng", 140), ("Storefront/API", 520), ("PostgreSQL", 930), ("Staff/Admin", 1350)]
    for name, x in actors:
        d.rounded_rectangle((x - 100, 100, x + 100, 165), radius=12, fill=PALE, outline=BLUE, width=3)
        d.text((x - 80, 118), name, font=font(20, True), fill=DARK)
        d.line((x, 165, x, 920), fill=MUTED, width=2)
    messages = [
        (140, 520, 235, "Checkout chuyển khoản"),
        (520, 930, 330, "Tạo Order + Payment(PENDING)"),
        (930, 520, 425, "orderNumber + payment reference"),
        (520, 140, 520, "Hiển thị hướng dẫn chuyển khoản"),
        (1350, 520, 650, "Xác nhận giao dịch đã nhận"),
        (520, 930, 745, "Conditional update Payment=PAID + AuditLog"),
        (930, 520, 840, "Kết quả đối soát"),
    ]
    for x1, x2, y, text in messages:
        arrow(d, (x1, y), (x2, y), GREEN if "PAID" in text else BLUE, 3)
        d.text((min(x1, x2) + 20, y - 32), text, font=font(18, True), fill=INK)
    image.save(path)
    return path


def diagram_sdlc():
    image, d, path = save_canvas("enterprise_sdlc.png", (1700, 950))
    d.text((55, 25), "Quy trình phát triển và vận hành dự án", font=font(38, True), fill=DARK)
    stages = [
        ("1. Khởi tạo", "Mục tiêu, phạm vi,\nstakeholder"),
        ("2. Discovery", "Yêu cầu, quy trình,\nrủi ro"),
        ("3. Thiết kế", "UX, kiến trúc,\ndữ liệu"),
        ("4. Xây dựng", "Frontend, backend,\nintegration"),
        ("5. Kiểm thử", "Unit, integration,\nE2E, UAT"),
        ("6. Phát hành", "CI/CD, migration,\nrollback"),
        ("7. Vận hành", "SLO, monitoring,\nincident"),
    ]
    x = 45
    positions = []
    for i, (title, detail) in enumerate(stages):
        y = 180 if i % 2 == 0 else 500
        box(d, (x, y, x + 205, y + 190), title, detail.split("\n"), LIGHT, BLUE, BLUE)
        positions.append((x + 102, y + 95))
        x += 235
    for i in range(len(positions) - 1):
        arrow(d, positions[i], positions[i + 1], GREEN, 4)
    arrow(d, (positions[-1][0], positions[-1][1] + 110), (positions[1][0], positions[1][1] + 110), GOLD, 4)
    d.text((540, 840), "Vòng phản hồi từ vận hành quay lại discovery và backlog", font=font(22, True), fill=GOLD)
    image.save(path)
    return path


def diagram_deployment():
    image, d, path = save_canvas("enterprise_deployment.png", (1600, 950))
    d.text((55, 25), "Kiến trúc triển khai mục tiêu", font=font(38, True), fill=DARK)
    box(d, (70, 180, 370, 390), "Người dùng", ["Web browser", "TLS", "Responsive UI"], PALE, BLUE, BLUE)
    box(d, (520, 130, 1060, 440), "Nền tảng ứng dụng", ["Next.js runtime", "Server/API routes", "Auth & authorization", "Rate limit / validation", "Structured logging"], LIGHT, BLUE, BLUE)
    box(d, (1210, 180, 1510, 390), "Dịch vụ ngoài", ["Cloudinary", "Resend", "Maps"], PALE, BLUE, BLUE)
    box(d, (260, 590, 650, 820), "PostgreSQL", ["Production database", "Migrations", "Backup + restore test"], LIGHT, GREEN, GREEN)
    box(d, (830, 590, 1220, 820), "Quan sát hệ thống", ["Error tracking", "Metrics & alert", "Audit log", "SLO dashboard"], LIGHT, GOLD, GOLD)
    arrow(d, (370, 285), (520, 285))
    arrow(d, (1060, 285), (1210, 285))
    arrow(d, (700, 440), (520, 590))
    arrow(d, (880, 440), (1020, 590))
    image.save(path)
    return path


def diagram_governance():
    image, d, path = save_canvas("enterprise_governance.png", (1600, 1000))
    d.text((55, 25), "Quality gates và quản trị thay đổi", font=font(38, True), fill=DARK)
    gates = [
        ("Backlog Ready", ["Acceptance criteria", "Risk & dependency"]),
        ("Design Ready", ["Architecture review", "Data & security"]),
        ("Code Ready", ["Review", "Lint/typecheck/test"]),
        ("Release Ready", ["UAT", "Migration/rollback"]),
        ("Operate Ready", ["SLO", "Runbook/alert"]),
    ]
    y = 130
    for i, (title, lines) in enumerate(gates):
        x = 120 + (i % 2) * 780
        if i == 4:
            x = 510
        box(d, (x, y, x + 580, y + 145), title, lines, LIGHT, GREEN if i >= 2 else BLUE, GREEN if i >= 2 else BLUE)
        if i % 2 == 1:
            y += 190
    d.text((210, 870), "Mỗi gate có chủ sở hữu, bằng chứng kiểm chứng và tiêu chí không được bỏ qua.", font=font(24, True), fill=DARK)
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


class Report:
    def __init__(self, title, subtitle, preset):
        self.doc = Document()
        self.title = title
        self.subtitle = subtitle
        self.preset = preset
        self.fig_no = 0
        self.table_no = 0
        self.setup()

    def setup(self):
        sec = self.doc.sections[0]
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = rgb(INK)
        normal.paragraph_format.space_after = Pt(8 if self.preset == "narrative" else 6)
        normal.paragraph_format.line_spacing = 1.333 if self.preset == "narrative" else 1.10
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.preset == "narrative" else WD_ALIGN_PARAGRAPH.LEFT
        for style_name, size, color, before, after in [
            ("Title", 28, DARK, 0, 12),
            ("Subtitle", 14, MUTED, 0, 10),
            ("Heading 1", 16, BLUE, 18 if self.preset == "narrative" else 16, 10 if self.preset == "narrative" else 8),
            ("Heading 2", 13, BLUE, 12, 6),
            ("Heading 3", 12, DARK, 8, 4),
        ]:
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = rgb(color)
            style.font.bold = style_name != "Subtitle"
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        header = sec.header.paragraphs[0]
        header.text = "MAISON DE FLOF  |  HỒ SƠ DỰ ÁN"
        header.style = styles["Caption"]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.color.rgb = rgb(MUTED)
        add_page_number(sec.footer.paragraphs[0])
        sec.footer.paragraphs[0].runs[0].font.color.rgb = rgb(MUTED)

    def cover(self, report_type, summary):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(110)
        r = p.add_run("MAISON DE FLOF")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = rgb(BLUE)
        p = self.doc.add_paragraph(self.title, style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(30)
        p = self.doc.add_paragraph(self.subtitle, style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(30)
        table = self.doc.add_table(rows=4, cols=2)
        values = [
            ("Loại tài liệu", report_type),
            ("Đối tượng hệ thống", "Website thương mại điện tử sơn và tư vấn màu Maison de FLOF"),
            ("Cơ sở phân tích", "Mã nguồn, Prisma schema, migration, test, CI, audit và roadmap trong repository"),
            ("Ngày lập báo cáo", "11/06/2026"),
        ]
        for row, (a, b) in zip(table.rows, values):
            row.cells[0].text = a
            row.cells[1].text = b
            set_cell_shading(row.cells[0], PALE)
            row.cells[0].paragraphs[0].runs[0].bold = True
        set_table_geometry(table, [2700, 6660])
        self.doc.add_paragraph("")
        p = self.doc.add_paragraph(summary)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(18)
        self.doc.add_paragraph("")
        p = self.doc.add_paragraph("Phiên bản 1.0 | Tài liệu nội bộ phục vụ đồ án và quản trị dự án")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = rgb(MUTED)
        self.doc.add_page_break()

    def h1(self, text):
        return self.doc.add_heading(text, level=1)

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def p(self, text, bold_prefix=None):
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            r.bold = True
            p.add_run(text[len(bold_prefix):])
        else:
            p.add_run(text)
        return p

    def bullets(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.add_run(item)
            p.paragraph_format.space_after = Pt(4 if self.preset == "narrative" else 8)

    def numbered(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Number")
            p.add_run(item)
            p.paragraph_format.space_after = Pt(4 if self.preset == "narrative" else 8)

    def callout(self, title, text, fill=PALE):
        table = self.doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [9360])
        set_cell_shading(table.cell(0, 0), fill)
        p = table.cell(0, 0).paragraphs[0]
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.color.rgb = rgb(DARK)
        p.add_run(text)

    def table(self, headers, rows, widths=None):
        self.table_no += 1
        if widths is None:
            widths = [9360 // len(headers)] * len(headers)
            widths[-1] += 9360 - sum(widths)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            set_cell_shading(cell, LIGHT if self.preset == "business" else "#F4F6F9")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = rgb(DARK)
        set_repeat_table_header(table.rows[0])
        for row_values in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row_values):
                cells[i].text = str(value)
        set_table_geometry(table, widths)
        cap = self.doc.add_paragraph(f"Bảng {self.table_no}. {headers[0]} và thông tin liên quan")
        cap.style = self.doc.styles["Caption"]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return table

    def figure(self, path, caption, width=6.4):
        self.fig_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = self.doc.add_paragraph(f"Hình {self.fig_no}. {caption}")
        cap.style = self.doc.styles["Caption"]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def contents(self, chapters):
        self.h1("Mục lục nội dung")
        for title, sections in chapters:
            p = self.doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.color.rgb = rgb(DARK)
            for section in sections:
                p = self.doc.add_paragraph(section)
                p.paragraph_format.left_indent = Inches(0.25)
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)


def academic_report(diagrams):
    r = Report(
        "BÁO CÁO ĐỒ ÁN WEBSITE",
        "Phân tích, thiết kế và xây dựng hệ thống Maison de FLOF",
        "narrative",
    )
    r.cover(
        "Báo cáo đồ án phân tích và thiết kế hệ thống",
        "Tài liệu trình bày toàn bộ quá trình khảo sát, phân tích yêu cầu, mô hình hóa, thiết kế, cài đặt, kiểm thử và đánh giá website Maison de FLOF. Báo cáo dựa trên hiện trạng mã nguồn tại ngày 11/06/2026 và phân biệt rõ chức năng đã có với hạng mục cần tiếp tục hoàn thiện.",
    )
    r.contents([
        ("Tóm tắt và phạm vi", ["Tóm tắt dự án", "Phương pháp thực hiện", "Thuật ngữ"]),
        ("Chương 1. Khảo sát và phân tích yêu cầu", ["Bài toán, mục tiêu, tác nhân, yêu cầu chức năng và phi chức năng"]),
        ("Chương 2. Phân tích hệ thống", ["Use Case, quy trình nghiệp vụ, Activity và Sequence"]),
        ("Chương 3. Thiết kế hệ thống", ["Kiến trúc, dữ liệu, API, bảo mật và UI/UX"]),
        ("Chương 4. Cài đặt và xây dựng", ["Công nghệ, module, giao dịch thương mại, quản trị"]),
        ("Chương 5. Kiểm thử, triển khai và đánh giá", ["Chiến lược kiểm thử, CI, rủi ro, kết luận"]),
        ("Phụ lục", ["Danh mục API, mô hình dữ liệu, checklist nghiệm thu"]),
    ])
    r.h1("Tóm tắt dự án")
    r.p("Maison de FLOF là website thương mại điện tử chuyên sản phẩm sơn, bảng màu và dịch vụ tư vấn màu. Hệ thống kết hợp trải nghiệm mua sắm cho khách hàng với cổng quản trị phục vụ đơn hàng, thanh toán, tồn kho, nội dung, tư vấn và chăm sóc khách hàng.")
    r.callout("Hiện trạng định lượng", "Repository có 31 trang giao diện, 40 API routes, 29 model Prisma, 11 enum, 6 migration và 11 test tự động. CI chạy cài đặt, Prisma generate, lint, typecheck, test, dependency audit mức high và build.")
    r.h2("Mục tiêu đồ án")
    r.bullets([
        "Xây dựng trải nghiệm tìm kiếm, lựa chọn, tính toán và mua sản phẩm sơn trực tuyến.",
        "Thiết kế dữ liệu đảm bảo tính toàn vẹn cho đơn hàng, thanh toán, tồn kho và lịch sử trạng thái.",
        "Cung cấp cổng quản trị có phân quyền, audit log và công cụ vận hành.",
        "Thiết lập nền tảng chất lượng có migration, test và CI để tiếp tục phát triển.",
    ])
    r.h2("Phương pháp thực hiện")
    r.p("Quy trình áp dụng theo vòng đời phát triển hệ thống: khởi tạo, khảo sát, phân tích yêu cầu, phân tích hệ thống, thiết kế, xây dựng, tích hợp, kiểm thử, triển khai và bảo trì. Dữ liệu đánh giá được thu thập bằng cách đọc mã nguồn, schema, migration, test, workflow CI, audit report và roadmap.")
    r.table(["Thuật ngữ", "Ý nghĩa"], [
        ("RBAC", "Kiểm soát truy cập theo vai trò và quyền."),
        ("Idempotency", "Cùng một yêu cầu lặp lại không tạo thêm hiệu ứng nghiệp vụ."),
        ("Payment lifecycle", "Vòng đời trạng thái thanh toán từ chờ xử lý đến trả tiền/hoàn tiền."),
        ("Inventory ledger", "Sổ giao dịch tăng giảm tồn kho có tham chiếu và lịch sử."),
        ("UAT", "Kiểm thử chấp nhận của người dùng/nghiệp vụ."),
        ("SLO", "Mục tiêu mức dịch vụ dùng để vận hành và cảnh báo."),
    ], [2100, 7260])

    r.h1("Chương 1. Khảo sát và phân tích yêu cầu")
    r.h2("1.1 Bài toán và bối cảnh")
    r.p("Kinh doanh sơn có đặc thù nhiều biến thể màu, loại bề mặt, dung tích, độ phủ và nhu cầu tư vấn. Khách hàng cần không chỉ một danh mục sản phẩm mà còn công cụ tính lượng sơn, hình dung màu, tìm đại lý, nhận tư vấn và theo dõi đơn hàng. Doanh nghiệp đồng thời cần kiểm soát giá, coupon, tồn kho, thanh toán, nội dung và hoạt động của nhân viên.")
    r.h2("1.2 Phạm vi hệ thống")
    r.table(["Trong phạm vi", "Ngoài phạm vi hoặc chưa hoàn thiện"], [
        ("Storefront, catalog, màu, giỏ hàng, checkout, tài khoản, yêu cầu báo giá, chat", "Ứng dụng di động native và hệ thống ERP độc lập"),
        ("COD và chuyển khoản với payment record, đối soát/hoàn tiền thủ công", "Webhook ngân hàng/cổng thanh toán tự động"),
        ("Cổng admin quản lý vận hành, RBAC và audit log", "SIEM, SSO doanh nghiệp và phê duyệt nhiều cấp"),
        ("Unit/policy tests và CI quality gate", "Bộ integration/E2E đầy đủ, staging và monitoring production"),
    ], [4680, 4680])
    r.h2("1.3 Tác nhân")
    r.figure(diagrams["usecase"], "Use Case tổng quan của hệ thống")
    r.table(["Tác nhân", "Mục tiêu chính", "Quyền điển hình"], [
        ("Khách chưa đăng nhập", "Khám phá sản phẩm, màu và nội dung", "Xem catalog, tính sơn, tìm đại lý, đăng ký"),
        ("Khách hàng", "Mua hàng và quản lý quan hệ với thương hiệu", "Checkout, địa chỉ, yêu thích, đơn hàng, đánh giá"),
        ("Staff", "Vận hành đơn, payment, kho và hỗ trợ", "Quyền giới hạn theo policy, không toàn quyền catalog/người dùng"),
        ("Admin", "Quản trị toàn hệ thống", "Catalog, tài khoản, cấu hình và audit"),
        ("Dịch vụ ngoài", "Cung cấp media, email và bản đồ", "Cloudinary, Resend, MapLibre"),
    ], [1600, 3650, 4110])
    r.h2("1.4 Yêu cầu chức năng")
    r.table(["Mã", "Nhóm", "Yêu cầu", "Trạng thái"], [
        ("FR-01", "Catalog", "Xem, lọc và tìm kiếm sản phẩm/màu", "Đã có"),
        ("FR-02", "Công cụ", "Tính lượng sơn và trực quan hóa màu", "Đã có"),
        ("FR-03", "Commerce", "Giỏ hàng, coupon, checkout, tạo đơn", "Đã có"),
        ("FR-04", "Payment", "COD/chuyển khoản, đối soát và hoàn tiền thủ công", "Đã có nền tảng"),
        ("FR-05", "Account", "Hồ sơ, địa chỉ, yêu thích, lịch sử đơn", "Đã có"),
        ("FR-06", "Support", "Chat, báo giá, dealer locator", "Đã có"),
        ("FR-07", "Admin", "Quản lý order/payment/inventory/catalog/content", "Đã có"),
        ("FR-08", "Governance", "RBAC chi tiết và audit log", "Đã có lõi"),
    ], [900, 1450, 4800, 2210])
    r.h2("1.5 Yêu cầu phi chức năng")
    r.table(["Nhóm", "Yêu cầu", "Cách đáp ứng hiện tại", "Khoảng trống"], [
        ("Bảo mật", "Xác thực, phân quyền, validation, bảo vệ mutation", "Auth.js, role/permission policy, Zod, audit", "Rate limit dùng chung, secret governance"),
        ("Tin cậy", "Không trùng đơn, không hoàn kho/cộng doanh thu hai lần", "Idempotency, conditional update, transaction", "Cần integration/concurrency tests"),
        ("Hiệu năng", "Trang phản hồi nhanh, truy vấn có kiểm soát", "Next.js, server API, Prisma", "Pagination và đo hiệu năng"),
        ("Khả dụng", "Responsive và tiếp cận được", "Responsive UI", "Cần WCAG 2.2 AA audit"),
        ("Vận hành", "Có CI, deploy, backup, monitoring, runbook", "CI hiện hữu", "Thiếu staging, deploy pipeline, SLO/monitoring"),
    ], [1400, 2500, 2900, 2560])

    r.h1("Chương 2. Phân tích hệ thống")
    r.h2("2.1 Phân rã miền nghiệp vụ")
    r.p("Hệ thống được phân tích thành các miền: danh mục và màu sắc; khách hàng và danh tính; giỏ hàng và checkout; đơn hàng, payment và tồn kho; nội dung và tư vấn; quản trị và governance. Cách phân rã này giúp mỗi invariant quan trọng có chủ sở hữu rõ ràng.")
    r.h2("2.2 Quy trình checkout")
    r.figure(diagrams["activity"], "Activity Diagram checkout có chống tạo đơn trùng")
    r.numbered([
        "Client gửi dữ liệu giỏ hàng và Idempotency-Key duy nhất.",
        "Server xác thực, chuẩn hóa và kiểm tra khóa đã được xử lý hay chưa.",
        "Giá, giảm giá, coupon và tồn kho được tính lại phía server; không tin giá từ client.",
        "Transaction tạo Order, OrderItem, Payment và inventory transaction tương ứng.",
        "Kết quả được gắn với khóa idempotency để retry trả lại cùng một đơn.",
    ])
    r.h2("2.3 Vòng đời đơn hàng và cạnh tranh đồng thời")
    r.p("Cập nhật trạng thái đơn sử dụng conditional update trong transaction. Chỉ request thỏa trạng thái nguồn mới được chuyển trạng thái; request đồng thời còn lại không được lặp lại hiệu ứng như hoàn kho hoặc cộng doanh thu. OrderStatusHistory lưu diễn tiến để phục vụ truy vết.")
    r.table(["Trạng thái", "Ý nghĩa", "Chuyển tiếp tiêu biểu", "Invariant"], [
        ("PENDING", "Đơn mới tạo", "CONFIRMED hoặc CANCELLED", "Chưa ghi nhận hoàn tất"),
        ("CONFIRMED", "Đã xác nhận xử lý", "PROCESSING/CANCELLED", "Không xác nhận lặp"),
        ("PROCESSING", "Đang chuẩn bị", "SHIPPED/CANCELLED", "Kiểm soát tồn kho"),
        ("SHIPPED", "Đã giao vận", "DELIVERED", "Theo dõi giao hàng"),
        ("DELIVERED", "Hoàn tất giao", "Kết thúc", "Doanh thu chỉ ghi nhận một lần"),
        ("CANCELLED", "Đã hủy", "Kết thúc", "Hoàn kho chỉ một lần"),
    ], [1500, 2300, 2600, 2960])
    r.h2("2.4 Quy trình thanh toán chuyển khoản")
    r.figure(diagrams["sequence"], "Sequence Diagram đối soát chuyển khoản thủ công")
    r.p("Payment lifecycle hiện hỗ trợ PENDING, PAID, FAILED, CANCELLED và REFUNDED. Nhân viên có quyền phù hợp xác nhận hoặc hoàn tiền thủ công, và các mutation quan trọng được ghi AuditLog. Đây là nền tảng đúng cho vận hành, nhưng chưa phải đối soát tự động với ngân hàng hoặc cổng thanh toán.")
    r.h2("2.5 Ma trận truy vết yêu cầu")
    r.table(["Yêu cầu", "Thành phần thực hiện", "Dữ liệu", "Bằng chứng kiểm chứng"], [
        ("Chống đơn trùng", "Checkout API", "CheckoutIdempotency", "Security-commerce tests"),
        ("Theo dõi payment", "Admin payments API/UI", "Payment", "Payment lifecycle policy"),
        ("Không cập nhật trạng thái lặp", "Order transition service/API", "OrderStatusHistory", "Conditional update tests"),
        ("Phân quyền staff", "Permissions policy + guards", "Role/User", "Policy tests"),
        ("Truy vết thao tác", "Audit helper + admin API", "AuditLog", "Mutation review"),
    ], [1900, 2700, 2100, 2660])

    r.h1("Chương 3. Thiết kế hệ thống")
    r.h2("3.1 Kiến trúc tổng thể")
    r.figure(diagrams["architecture"], "Kiến trúc logic của Maison de FLOF")
    r.p("Ứng dụng sử dụng mô hình full-stack Next.js App Router. Lớp giao diện và API cùng repository, nhưng các trách nhiệm nghiệp vụ quan trọng được tổ chức qua route handlers, policy và helper. Prisma là lớp truy cập dữ liệu, PostgreSQL là nguồn dữ liệu chính, còn Cloudinary, Resend và MapLibre phục vụ media, email và bản đồ.")
    r.h2("3.2 Thiết kế dữ liệu")
    r.figure(diagrams["erd"], "ERD rút gọn theo các miền nghiệp vụ")
    r.p("Schema có 29 model và 11 enum. Thiết kế sử dụng snapshot thông tin đơn hàng để lịch sử không bị thay đổi khi catalog hoặc địa chỉ khách hàng thay đổi. Payment, inventory transaction, status history, idempotency và audit log tạo thành lớp bảo vệ tính toàn vẹn giao dịch.")
    r.table(["Miền", "Model tiêu biểu", "Quy tắc thiết kế"], [
        ("Danh tính", "Role, User, Account, Session, Customer, Address", "Tách xác thực, hồ sơ và địa chỉ"),
        ("Catalog", "Category, Supplier, Paint, PaintColor, Collection", "Quan hệ màu/sơn và metadata sản phẩm"),
        ("Commerce", "Order, OrderItem, Coupon, Payment", "Snapshot và transaction"),
        ("Kho", "InventoryTransaction", "Ledger thay cho sửa tồn trực tiếp"),
        ("Governance", "CheckoutIdempotency, AuditLog, OrderStatusHistory", "Chống lặp và truy vết"),
        ("Engagement", "Wishlist, Review, Blog, ChatMessage, QuoteRequest", "Tương tác và nội dung"),
    ], [1500, 3600, 4260])
    r.h2("3.3 Thiết kế API")
    r.p("Repository có 40 API routes, chia thành API công khai/khách hàng và API quản trị. Mọi mutation nhạy cảm cần đi qua xác thực, authorization, validation, transaction khi chạm tiền/kho/trạng thái, và audit log khi do staff/admin thực hiện.")
    r.table(["Lớp API", "Ví dụ", "Nguyên tắc"], [
        ("Công khai", "/api/products, /api/colors, /api/blog, /api/dealers", "Chỉ đọc; cần pagination/filter phía server khi dữ liệu lớn"),
        ("Khách hàng", "/api/orders, /api/profile/*, /api/chat", "Xác thực ownership và validate input"),
        ("Quản trị", "/api/admin/orders, payments, inventory, audit-logs", "Permission policy chi tiết và audit"),
        ("Tích hợp", "Cloudinary, Resend, MapLibre", "Tách secret, timeout, retry và xử lý lỗi"),
    ], [1800, 3600, 3960])
    r.h2("3.4 Thiết kế bảo mật")
    r.bullets([
        "Auth.js/NextAuth kết hợp Prisma adapter quản lý xác thực và session.",
        "RBAC được cụ thể hóa thành permission policy; Staff không mặc nhiên có toàn quyền.",
        "Zod validation bảo vệ dữ liệu đầu vào; giá và tổng tiền được tính lại phía server.",
        "AuditLog ghi hành động quan trọng; dependency audit được chạy trong CI ở mức high.",
        "Cần tiếp tục bổ sung rate limiting dùng chung, monitoring bảo mật, rotation secret và kiểm thử authorization toàn diện.",
    ])
    r.h2("3.5 Thiết kế UI/UX")
    r.p("Thiết kế hướng đến thương hiệu sơn cao cấp, tập trung hình ảnh màu sắc, công cụ hỗ trợ quyết định và luồng mua hàng rõ ràng. Các trang trọng yếu gồm home, catalog, chi tiết sản phẩm, màu, visualizer, calculator, cart, checkout, profile và admin. Khoảng trống cần xử lý là accessibility audit, cỡ chữ rất nhỏ tại một số vị trí và kiểm thử responsive có hệ thống.")

    r.h1("Chương 4. Cài đặt và xây dựng")
    r.h2("4.1 Công nghệ sử dụng")
    r.table(["Lớp", "Công nghệ", "Vai trò"], [
        ("Frontend", "Next.js 15, React 19, TypeScript, Tailwind, Framer Motion", "Giao diện, routing và tương tác"),
        ("State/Form", "Zustand, React Hook Form, Zod, TanStack Query", "Trạng thái, form, validation và dữ liệu"),
        ("Backend", "Next.js route handlers, Auth.js", "API, xác thực và nghiệp vụ"),
        ("Data", "Prisma 6, PostgreSQL/Neon", "ORM, migration và lưu trữ"),
        ("External", "Cloudinary, Resend, MapLibre", "Media, email và bản đồ"),
        ("Quality", "ESLint, TypeScript, Node test runner, GitHub Actions", "Quality gates"),
    ], [1600, 4100, 3660])
    r.h2("4.2 Các module đã xây dựng")
    r.bullets([
        "Storefront và catalog: sản phẩm, màu, collection, blog, dealer.",
        "Công cụ quyết định mua: color visualizer, paint calculator, wishlist và review.",
        "Commerce: cart, coupon, checkout, order, payment lifecycle và inventory ledger.",
        "Customer service: profile, địa chỉ, order history, quote request và chat.",
        "Admin: dashboard, order/invoice/payment, inventory, catalog, content, support, account và audit log.",
    ])
    r.h2("4.3 Kiểm soát tính toàn vẹn giao dịch")
    r.p("Các cải tiến sau audit tập trung vào invariant thương mại: thêm Payment và trạng thái; thêm CheckoutIdempotency; sửa race condition chuyển trạng thái đơn; chặn sửa stock trực tiếp; thu hẹp quyền Staff; thêm AuditLog; nâng dependency có advisory và tăng test từ 6 lên 11 trường hợp.")
    r.callout("Nguyên tắc triển khai", "Mutation liên quan tiền, kho hoặc trạng thái phải có transaction/idempotency/concurrency control phù hợp. Mutation do staff/admin thực hiện phải có authorization cụ thể và audit log.", "#EAF7F5")
    r.h2("4.4 Cấu trúc triển khai")
    r.table(["Thành phần", "Vị trí", "Nội dung"], [
        ("Trang và API", "src/app", "31 page.tsx và 40 route.ts"),
        ("Thành phần dùng chung", "src/components, src/lib", "UI, policy, nghiệp vụ và tích hợp"),
        ("Dữ liệu", "prisma/schema.prisma, prisma/migrations", "29 model và 6 migration"),
        ("Kiểm thử", "tests/*.test.ts", "Commerce, security-commerce, paint calculator"),
        ("CI", ".github/workflows/ci.yml", "Lint, typecheck, test, audit, build"),
    ], [1800, 3100, 4460])

    r.h1("Chương 5. Kiểm thử, triển khai và đánh giá")
    r.h2("5.1 Chiến lược kiểm thử")
    r.table(["Tầng", "Mục tiêu", "Hiện trạng", "Hành động tiếp theo"], [
        ("Static", "Lỗi cú pháp, type và style", "Lint + typecheck trong CI", "Duy trì bắt buộc"),
        ("Unit/Policy", "Logic tính toán, transition, permission", "11 trường hợp tự động", "Mở rộng edge cases"),
        ("Integration", "API + database + transaction", "Chưa đầy đủ", "Test checkout/payment/order thật"),
        ("E2E", "Luồng người dùng từ UI đến DB", "Chưa có bộ hoàn chỉnh", "Playwright cho luồng trọng yếu"),
        ("UAT/Non-functional", "Nghiệp vụ, accessibility, tải, bảo mật", "Chưa chuẩn hóa", "Checklist và môi trường staging"),
    ], [1400, 2500, 2300, 3160])
    r.h2("5.2 CI và triển khai")
    r.p("GitHub Actions hiện chạy trên pull request và push main với các bước npm ci, Prisma generate, lint, typecheck, test, dependency audit mức high và build. Đây là quality gate tốt ở cấp repository. Để production-ready cần thêm preview deployment, staging riêng, migration/rollback có kiểm soát, backup/restore test, error tracking, metrics, alert và runbook.")
    r.h2("5.3 Đánh giá rủi ro còn lại")
    r.table(["Rủi ro", "Mức", "Tác động", "Kiểm soát đề xuất"], [
        ("Chuyển khoản chưa đối soát tự động", "Cao", "Sai lệch/độ trễ vận hành", "Webhook/provider hoặc quy trình đối soát có SLA"),
        ("Thiếu integration/E2E", "Cao", "Regression luồng mua hàng", "Ưu tiên checkout, payment, order transition"),
        ("Thiếu monitoring/runbook", "Cao", "Phát hiện và khôi phục chậm", "SLO, alert, incident và DR"),
        ("2 advisory moderate gián tiếp", "Trung bình", "Rủi ro dependency", "Theo dõi bản vá tương thích, không force downgrade"),
        ("Pagination/accessibility chưa đồng đều", "Trung bình", "Hiệu năng và trải nghiệm", "Audit và backlog có tiêu chí đo"),
    ], [2700, 1200, 2600, 2860])
    r.h2("5.4 Kết quả và hạn chế")
    r.p("Dự án đã vượt mức prototype nhờ có schema tương đối đầy đủ, transaction, snapshot, migration, role guard, CI và các kiểm soát commerce mới. Tuy nhiên, mức sẵn sàng production vẫn phụ thuộc vào bằng chứng integration/E2E, đối soát thanh toán thật, staging, monitoring, backup và runbook. Vì vậy kết luận phù hợp là hệ thống có nền tảng tốt và có thể tiếp tục hardening, chưa nên coi là vận hành production hoàn chỉnh.")
    r.h2("5.5 Hướng phát triển")
    r.numbered([
        "Hoàn thiện payment provider/webhook hoặc quy trình đối soát chuẩn hóa có SLA.",
        "Xây bộ integration và E2E cho checkout, payment, hủy/hoàn tất đơn và phân quyền.",
        "Thiết lập staging, deploy pipeline, migration/rollback và backup/restore test.",
        "Bổ sung logging tập trung, error tracking, metrics, alert, SLO và runbook.",
        "Thực hiện performance, accessibility WCAG 2.2 AA và security review định kỳ.",
    ])
    r.h1("Kết luận")
    r.p("Maison de FLOF cho thấy một quy trình phát triển website hoàn chỉnh không chỉ dừng ở giao diện và tính năng, mà phải nối liền yêu cầu, thiết kế dữ liệu, invariant giao dịch, bảo mật, kiểm thử và vận hành. Báo cáo xác định rõ phần đã xây dựng và lộ trình để hệ thống tiến tới mức production-ready có thể kiểm chứng.")

    r.h1("Phụ lục A. Danh mục API theo miền")
    r.table(["Miền", "API tiêu biểu"], [
        ("Catalog", "/api/products, /api/colors, /api/categories, /api/color-collections, /api/dealers"),
        ("Customer", "/api/profile, /api/profile/addresses, /api/profile/favorites, /api/profile/password"),
        ("Commerce", "/api/orders, /api/orders/[orderNumber], /api/coupons/validate"),
        ("Support/Content", "/api/chat, /api/quote-request, /api/blog, /api/reviews"),
        ("Admin", "/api/admin/orders, payments, inventory, products, coupons, users, audit-logs và các route quản trị khác"),
    ], [2100, 7260])
    r.h1("Phụ lục B. Checklist nghiệm thu kỹ thuật")
    r.bullets([
        "Mọi yêu cầu có acceptance criteria và bằng chứng kiểm chứng.",
        "Mọi mutation nhạy cảm có authentication, authorization và validation.",
        "Tiền, kho, coupon và trạng thái có transaction/idempotency/concurrency control.",
        "Staff/Admin mutation quan trọng tạo audit log.",
        "Migration có kế hoạch deploy/rollback và được thử trên staging.",
        "Lint, typecheck, test, audit và build đạt trước khi phát hành.",
        "Luồng trọng yếu có integration/E2E và UAT.",
        "Production có monitoring, backup, alert, SLO và runbook.",
    ])
    path = OUT / "BAO_CAO_DO_AN_MAISON_DE_FLOF.docx"
    r.save(path)
    return path


def enterprise_report(diagrams):
    r = Report(
        "BÁO CÁO DOANH NGHIỆP",
        "Quy trình xây dựng, quản trị và vận hành dự án Maison de FLOF",
        "business",
    )
    r.cover(
        "Báo cáo quản trị vòng đời dự án website",
        "Tài liệu chuyển hóa hiện trạng kỹ thuật của Maison de FLOF thành quy trình doanh nghiệp có mục tiêu, vai trò, quality gate, kiểm soát rủi ro, kế hoạch phát hành và vận hành. Báo cáo dùng làm hồ sơ quản trị, nghiệm thu và định hướng production readiness.",
    )
    r.contents([
        ("Tóm tắt điều hành", ["Hiện trạng, kết luận và quyết định ưu tiên"]),
        ("1. Project charter và business case", ["Mục tiêu, phạm vi, stakeholder, KPI"]),
        ("2. Quy trình delivery end-to-end", ["Khởi tạo đến vận hành"]),
        ("3. Quản trị yêu cầu, kiến trúc và dữ liệu", ["Artifacts, design review và governance"]),
        ("4. Quản trị chất lượng, bảo mật và phát hành", ["Test, CI/CD, security và release gates"]),
        ("5. Vận hành, rủi ro và roadmap", ["SLO, incident, risk register và lộ trình"]),
        ("Phụ lục", ["RACI, Definition of Done, checklist release"]),
    ])
    r.h1("Tóm tắt điều hành")
    r.p("Maison de FLOF đã có phạm vi chức năng rộng và nền tảng kỹ thuật tốt hơn một prototype thông thường. Hệ thống có storefront, cổng quản trị, schema 29 model, migration, CI, transaction, idempotency, payment record, inventory ledger, RBAC và audit log. Các khoảng trống quyết định mức sẵn sàng production hiện nằm ở integration/E2E, đối soát thanh toán thật, deploy/staging, monitoring, backup và runbook.")
    r.callout("Khuyến nghị điều hành", "Tạm ưu tiên hardening và vận hành thay vì mở rộng tính năng. Release production chỉ nên được phê duyệt khi checkout/payment/order transition có bằng chứng E2E, migration và rollback được thử, monitoring hoạt động, và trách nhiệm vận hành đã được giao rõ.", "#EAF7F5")
    r.table(["Chỉ số hiện trạng", "Giá trị", "Ý nghĩa quản trị"], [
        ("Trang giao diện", "31", "Phạm vi trải nghiệm đã tương đối rộng"),
        ("API routes", "40", "Cần contract/authorization test có hệ thống"),
        ("Model / enum", "29 / 11", "Miền dữ liệu đủ lớn để cần governance"),
        ("Migration", "6", "Đã có nền tảng quản lý thay đổi DB"),
        ("Test tự động", "11", "Tăng từ 6 nhưng chưa đủ integration/E2E"),
        ("Dependency audit", "2 moderate còn lại", "Theo dõi bản vá tương thích; không force downgrade"),
    ], [2500, 1900, 4960])

    r.doc.add_page_break()
    r.h1("1. Project charter và business case")
    r.h2("1.1 Business case")
    r.p("Dự án tạo một kênh bán hàng và tư vấn màu trực tuyến giúp khách hàng đưa ra quyết định mua sơn tốt hơn, đồng thời chuẩn hóa vận hành đơn hàng, payment, tồn kho, nội dung và chăm sóc khách hàng. Giá trị doanh nghiệp đến từ tăng chuyển đổi, giảm thao tác thủ công, cải thiện khả năng truy vết và tạo nền tảng dữ liệu cho quản trị.")
    r.h2("1.2 Mục tiêu")
    r.table(["Mục tiêu", "Kết quả mong đợi", "KPI đề xuất"], [
        ("Doanh thu số", "Khách có thể khám phá đến checkout trọn vẹn", "Conversion, cart abandonment, order success"),
        ("Hiệu quả vận hành", "Đơn/payment/kho được xử lý có lịch sử", "Thời gian xử lý, sai lệch đối soát"),
        ("Quản trị rủi ro", "Thao tác nhạy cảm được phân quyền và audit", "Unauthorized attempts, audit coverage"),
        ("Chất lượng", "Mỗi release có bằng chứng kiểm chứng", "Pass rate, escaped defects, change failure rate"),
        ("Độ tin cậy", "Phát hiện và khôi phục sự cố nhanh", "Availability, MTTD, MTTR"),
    ], [2100, 4200, 3060])
    r.h2("1.3 Phạm vi và nguyên tắc ưu tiên")
    r.bullets([
        "Phạm vi cốt lõi: storefront, commerce, account, support, content và admin.",
        "Bắt buộc trước production: tính toàn vẹn giao dịch, payment control, authorization, test, migration, monitoring và runbook.",
        "Có thể sau bản đầu: automation nâng cao, analytics chuyên sâu, ERP/CRM integration và mobile native.",
        "Không mở rộng tính năng nếu quality gate của luồng commerce trọng yếu chưa đạt.",
    ])
    r.h2("1.4 Stakeholder và RACI")
    r.table(["Hoạt động", "Business owner", "PM/BA", "Tech lead/Dev", "QA", "Operations"], [
        ("Phê duyệt phạm vi/KPI", "A", "R", "C", "C", "C"),
        ("Yêu cầu và acceptance criteria", "A", "R", "C", "C", "I"),
        ("Kiến trúc và dữ liệu", "I", "C", "A/R", "C", "C"),
        ("Kiểm thử và UAT", "C", "C", "R", "A/R", "C"),
        ("Release và migration", "I", "C", "R", "C", "A/R"),
        ("Incident và hậu kiểm", "I", "C", "R", "C", "A/R"),
    ], [2500, 1372, 1372, 1372, 1372, 1372])
    r.p("Quy ước: A = Accountable, R = Responsible, C = Consulted, I = Informed.")

    r.h1("2. Quy trình delivery end-to-end")
    r.figure(diagrams["sdlc"], "Quy trình phát triển và vận hành dự án")
    r.h2("2.1 Khởi tạo")
    r.p("Xác định business case, sponsor, phạm vi, KPI, ngân sách, giả định và ràng buộc. Đầu ra bắt buộc gồm project charter, stakeholder map, high-level roadmap và cơ chế ra quyết định.")
    r.h2("2.2 Discovery và phân tích yêu cầu")
    r.p("Khảo sát hành vi khách hàng, quy trình nhân viên, dữ liệu hiện hữu và rủi ro vận hành. Mỗi requirement phải có chủ sở hữu, acceptance criteria, độ ưu tiên, dependency và traceability đến thiết kế/test.")
    r.h2("2.3 Thiết kế")
    r.p("Thiết kế bao gồm user flow, wireframe/UI, kiến trúc logic, ERD, API contract, security model và kế hoạch migration. Design review tập trung vào invariant tiền/kho/trạng thái, quyền hạn, khả năng vận hành và chi phí thay đổi.")
    r.h2("2.4 Lập kế hoạch và xây dựng")
    r.p("Backlog được chia theo vertical slice có thể kiểm chứng. Mỗi thay đổi đi qua code review, lint, typecheck, test và build. Các feature flag hoặc chiến lược rollout được dùng cho thay đổi có rủi ro cao.")
    r.h2("2.5 Kiểm thử, phát hành và vận hành")
    r.p("Bộ kiểm thử bao gồm unit, policy, integration, E2E, UAT và non-functional theo mức rủi ro. Release cần kế hoạch migration/rollback, observability, runbook và người trực vận hành. Sau phát hành phải review KPI, lỗi và bài học để cập nhật backlog.")
    r.table(["Giai đoạn", "Artifact bắt buộc", "Gate phê duyệt"], [
        ("Khởi tạo", "Charter, scope, KPI, stakeholder", "Sponsor phê duyệt giá trị và phạm vi"),
        ("Discovery", "Requirement, process map, risk log", "Backlog ready"),
        ("Thiết kế", "Architecture, ERD, API, UX, threat model", "Design ready"),
        ("Xây dựng", "Code, migration, test, documentation", "Code ready"),
        ("Phát hành", "UAT, release plan, rollback, runbook", "Release ready"),
        ("Vận hành", "Dashboard, alert, incident record, review", "Operate ready"),
    ], [1700, 4400, 3260])

    r.h1("3. Quản trị yêu cầu, kiến trúc và dữ liệu")
    r.h2("3.1 Quản trị yêu cầu")
    r.table(["Trường requirement", "Nội dung tối thiểu"], [
        ("Mã và mô tả", "Một mục tiêu nghiệp vụ rõ, không trộn nhiều hành vi"),
        ("Actor và trigger", "Ai khởi tạo và sự kiện nào bắt đầu"),
        ("Luồng chính/ngoại lệ", "Happy path, validation, retry, timeout và lỗi"),
        ("Acceptance criteria", "Điều kiện có thể kiểm thử và đo lường"),
        ("Dữ liệu/quyền", "Dữ liệu đọc/ghi, ownership và permission"),
        ("Bằng chứng", "Liên kết thiết kế, code, test và release"),
    ], [2600, 6760])
    r.h2("3.2 Kiến trúc hiện tại và mục tiêu")
    r.figure(diagrams["deployment"], "Kiến trúc triển khai mục tiêu")
    r.p("Kiến trúc hiện tại là Next.js full-stack với PostgreSQL/Prisma và dịch vụ ngoài. Để vận hành doanh nghiệp, cần bổ sung lớp quan sát hệ thống, môi trường tách biệt, backup/restore, secret governance, rate limit dùng chung và quy trình deploy có rollback.")
    r.h2("3.3 Governance dữ liệu")
    r.table(["Dữ liệu", "Owner", "Kiểm soát bắt buộc", "Retention/đánh giá"], [
        ("Tài khoản và hồ sơ", "Product/Operations", "Least privilege, ownership, bảo vệ PII", "Rà soát truy cập định kỳ"),
        ("Đơn hàng/payment", "Finance/Operations", "Transaction, idempotency, audit, reconciliation", "Theo chính sách tài chính"),
        ("Tồn kho", "Operations", "Ledger, reference, không sửa trực tiếp", "Đối chiếu định kỳ"),
        ("Audit log", "Security/Operations", "Immutable, truy cập giới hạn", "Theo nhu cầu điều tra/compliance"),
        ("Media/nội dung", "Marketing", "Ownership, approval, xóa có audit", "Lifecycle policy"),
    ], [2100, 1900, 3300, 2060])
    r.h2("3.4 Kiểm soát thương mại trọng yếu")
    r.bullets([
        "Server là nguồn tính giá, giảm giá và tổng tiền; không tin dữ liệu giá từ client.",
        "Checkout sử dụng Idempotency-Key để retry không tạo đơn trùng.",
        "Order transition dùng conditional update trong transaction để tránh race condition.",
        "Payment có trạng thái và lịch sử; đối soát/hoàn tiền thủ công có audit.",
        "Inventory dùng ledger và tham chiếu nghiệp vụ; không cho chỉnh stock trực tiếp.",
    ])

    r.h1("4. Quản trị chất lượng, bảo mật và phát hành")
    r.h2("4.1 Quality gates")
    r.figure(diagrams["governance"], "Quality gates và quản trị thay đổi")
    r.table(["Gate", "Tiêu chí không được thiếu", "Bằng chứng"], [
        ("Backlog Ready", "Acceptance criteria, priority, dependency, risk", "Requirement record"),
        ("Design Ready", "Architecture/data/security review", "Decision record, diagrams"),
        ("Code Ready", "Review, lint, typecheck, unit/policy test", "CI pass"),
        ("Release Ready", "Integration/E2E, UAT, migration và rollback", "Release checklist"),
        ("Operate Ready", "Monitoring, alert, on-call, runbook, backup", "Dashboard và rehearsal"),
    ], [1800, 4500, 3060])
    r.h2("4.2 Chiến lược kiểm thử theo rủi ro")
    r.p("Hiện tại 11 test tự động đã bao phủ một phần logic commerce, security-commerce và paint calculator. Do luồng payment/order/inventory có rủi ro tài chính, doanh nghiệp cần ưu tiên integration test với database thật và E2E cho retry checkout, xác nhận/hoàn tiền, hủy/hoàn tất đơn, permission và audit.")
    r.table(["Ưu tiên", "Kịch bản", "Loại test"], [
        ("P0", "Retry checkout cùng key chỉ tạo một order", "Integration + concurrency"),
        ("P0", "Hai request hủy/hoàn tất đồng thời không lặp side effect", "Integration + concurrency"),
        ("P0", "Staff không truy cập chức năng vượt quyền", "API authorization + E2E"),
        ("P0", "Payment PAID/REFUNDED có audit và không cập nhật lặp", "Integration"),
        ("P1", "Catalog, coupon, stock và checkout end-to-end", "E2E"),
        ("P1", "Responsive, accessibility và performance", "Non-functional"),
    ], [1200, 5700, 2460])
    r.h2("4.3 Bảo mật")
    r.table(["Lĩnh vực", "Hiện trạng", "Hành động doanh nghiệp"], [
        ("Identity/access", "Auth.js, role guard, permission policy", "Review quyền định kỳ, test deny-by-default"),
        ("Input/API", "Zod validation ở nhiều route", "Contract test, rate limit shared, abuse monitoring"),
        ("Audit", "AuditLog cho hành động quan trọng", "Retention, alert và truy vấn điều tra"),
        ("Dependency", "CI audit mức high; còn 2 moderate gián tiếp", "Theo dõi bản vá tương thích, exception có hạn dùng"),
        ("Secrets/data", "Env-based configuration", "Vault/rotation, phân tách môi trường, PII policy"),
    ], [2000, 3500, 3860])
    r.h2("4.4 CI/CD và release")
    r.p("CI hiện đã bảo vệ chất lượng code. CD mục tiêu cần tạo preview cho pull request, staging tách biệt, production deploy có approval, migration trước/đồng bộ với ứng dụng, health check, canary hoặc rollback nhanh. Không release khi thiếu bằng chứng cho các gate P0.")

    r.h1("5. Vận hành, rủi ro và roadmap")
    r.h2("5.1 Mô hình vận hành")
    r.table(["Năng lực", "Yêu cầu tối thiểu", "Chỉ số"], [
        ("Observability", "Structured logs, error tracking, metrics, trace theo request/order", "Error rate, latency, saturation"),
        ("Reliability", "Health check, backup/restore, rollback", "Availability, RPO, RTO"),
        ("Incident", "On-call, severity, communication, postmortem", "MTTD, MTTR, recurrence"),
        ("Commerce ops", "Payment reconciliation, order exception queue", "Sai lệch, aging, xử lý đúng SLA"),
        ("Security ops", "Audit review, dependency/secret process", "Findings và thời gian khắc phục"),
    ], [2000, 4400, 2960])
    r.h2("5.2 SLO đề xuất")
    r.table(["Dịch vụ", "SLO ban đầu", "Cảnh báo"], [
        ("Storefront/API", "99,5% khả dụng theo tháng", "Burn-rate và lỗi 5xx"),
        ("Checkout", "99% request hợp lệ tạo/khôi phục kết quả đúng", "Tăng lỗi, timeout, đơn trùng"),
        ("Payment operations", "Đối soát thủ công trong SLA đã cam kết", "Payment PENDING quá hạn"),
        ("Order operations", "Không lặp side effect do transition đồng thời", "Bất thường inventory/audit"),
        ("Backup", "Khôi phục thử nghiệm theo lịch", "Backup lỗi hoặc restore thất bại"),
    ], [2500, 4300, 2560])
    r.h2("5.3 Risk register")
    r.table(["ID", "Rủi ro", "Mức", "Owner", "Đối sách/tiêu chí đóng"], [
        ("R1", "Chưa có đối soát thanh toán tự động", "Cao", "Finance/Ops", "Provider/webhook hoặc SLA đối soát được đo"),
        ("R2", "Thiếu integration/E2E", "Cao", "Tech/QA", "Các kịch bản P0 chạy ổn định trong CI/staging"),
        ("R3", "Thiếu monitoring/runbook", "Cao", "Ops", "Dashboard, alert và incident rehearsal"),
        ("R4", "Migration chưa áp dụng/kiểm chứng mọi môi trường", "Cao", "Tech/Ops", "Staging deploy + rollback rehearsal"),
        ("R5", "2 advisory moderate gián tiếp", "Trung bình", "Tech", "Theo dõi patch, exception có ngày hết hạn"),
        ("R6", "Accessibility/pagination chưa đồng đều", "Trung bình", "Product/Tech", "Audit và tiêu chí đo đạt"),
    ], [700, 2800, 1100, 1600, 3160])
    r.h2("5.4 Roadmap đề xuất")
    r.table(["Giai đoạn", "Thời lượng tham chiếu", "Kết quả"], [
        ("P0 - Commerce hardening", "2-3 tuần", "Integration/concurrency tests, payment/reconciliation, permission review"),
        ("P1 - Release readiness", "2-3 tuần", "Staging, deploy/migration/rollback, monitoring, backup, runbook"),
        ("P2 - UAT và production launch", "1-2 tuần", "UAT, security/accessibility checks, rollout và hypercare"),
        ("P3 - Tối ưu vận hành", "Liên tục", "SLO review, automation, analytics và backlog cải tiến"),
    ], [2600, 2500, 4260])
    r.h2("5.5 Điều kiện phê duyệt production")
    r.numbered([
        "Mọi kịch bản P0 có integration/E2E và chạy đạt trên staging.",
        "Migration, rollback và backup/restore đã được diễn tập.",
        "Payment reconciliation có owner, SLA, dashboard/queue và audit.",
        "Monitoring, alert, incident process và runbook đã hoạt động.",
        "Không còn lỗ hổng high/critical; moderate có quyết định chấp nhận rủi ro có hạn dùng.",
        "Business owner hoàn tất UAT và ký release checklist.",
    ])
    r.h1("Kết luận điều hành")
    r.p("Maison de FLOF đã có nền tảng để trở thành hệ thống doanh nghiệp, nhưng năng lực production không được chứng minh bằng số lượng màn hình. Quyết định phát hành phải dựa trên invariant thương mại, bằng chứng kiểm thử, quyền hạn, khả năng quan sát và năng lực khôi phục. Quy trình trong báo cáo này biến các yêu cầu đó thành gate, trách nhiệm và tiêu chí nghiệm thu cụ thể.")

    r.doc.add_page_break()
    r.h1("Phụ lục A. Definition of Done")
    r.bullets([
        "Yêu cầu và acceptance criteria đã được cập nhật; phạm vi ngoài yêu cầu được ghi nhận.",
        "Thiết kế/API/data/security impact đã được review khi cần.",
        "Code review hoàn tất; không còn lỗi lint/typecheck/build.",
        "Test phù hợp mức rủi ro đã thêm và chạy đạt.",
        "Migration tương thích, rollback và dữ liệu mẫu đã được kiểm chứng.",
        "Authorization, audit và observability được bổ sung cho mutation nhạy cảm.",
        "Tài liệu vận hành/release note/runbook được cập nhật.",
    ])
    r.h1("Phụ lục B. Release checklist")
    r.table(["Hạng mục", "Câu hỏi kiểm tra", "Chủ sở hữu"], [
        ("Scope", "Các thay đổi và rủi ro đã được phê duyệt?", "PM/Business"),
        ("Quality", "CI, integration/E2E và UAT đều đạt?", "Tech/QA"),
        ("Database", "Migration/rollback/backup đã được thử?", "Tech/Ops"),
        ("Security", "Quyền, secret, dependency và audit đạt?", "Tech/Security"),
        ("Operations", "Dashboard, alert, on-call và runbook sẵn sàng?", "Ops"),
        ("Communication", "Kế hoạch rollout, support và thông báo rõ?", "PM/Ops"),
    ], [1800, 5200, 2360])
    path = OUT / "BAO_CAO_DOANH_NGHIEP_QUY_TRINH_DU_AN_MAISON_DE_FLOF.docx"
    r.save(path)
    return path


def main():
    diagrams = {
        "architecture": diagram_architecture(),
        "usecase": diagram_usecase(),
        "erd": diagram_erd(),
        "activity": diagram_activity(),
        "sequence": diagram_sequence(),
        "sdlc": diagram_sdlc(),
        "deployment": diagram_deployment(),
        "governance": diagram_governance(),
    }
    academic = academic_report(diagrams)
    enterprise = enterprise_report(diagrams)
    print(academic)
    print(enterprise)


if __name__ == "__main__":
    main()
